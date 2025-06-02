import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph as DocxParagraphObject
from docx.oxml.text.paragraph import CT_P
# from docx.text.paragraph import Paragraph as DocxParagraph # Já importado como DocxParagraphObject
# from docx.oxml.table import CT_Tbl # Não usado diretamente aqui

from datetime import datetime, timedelta
from num2words import num2words
from dateutil.relativedelta import relativedelta

# --- Funções de Manipulação de Documento DOCX (Baseadas em doc_utils.py do GitHub) ---
def replace_placeholder_in_paragraph(paragraph: DocxParagraphObject, placeholder: str, new_text_val: str, is_bold: bool = False): #
    str_new_text_val = str(new_text_val)
    current_paragraph_text = "".join(run.text for run in paragraph.runs)

    if placeholder not in current_paragraph_text:
        return

    temp_paragraph_text_parts = []
    last_idx = 0
    while placeholder in current_paragraph_text[last_idx:]:
        start_idx = current_paragraph_text.find(placeholder, last_idx)
        if start_idx == -1:
            break
        
        if start_idx > last_idx:
            temp_paragraph_text_parts.append( (current_paragraph_text[last_idx:start_idx], False) )
        
        temp_paragraph_text_parts.append( (str_new_text_val, is_bold) )
        last_idx = start_idx + len(placeholder)

    if last_idx < len(current_paragraph_text):
        temp_paragraph_text_parts.append( (current_paragraph_text[last_idx:], False) )

    for _ in range(len(paragraph.runs)):
        if paragraph.runs:
            p_element = paragraph._p
            p_element.remove(paragraph.runs[0]._r)

    for text_content, should_be_bold_segment in temp_paragraph_text_parts:
        if not text_content:
            continue
        run = paragraph.add_run(text_content)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run.font.bold = should_be_bold_segment

def replace_placeholders_in_document(doc: Document, replacements: dict, bold_keys: set = None) -> Document: #
    if bold_keys is None:
        bold_keys = set()

    for paragraph in doc.paragraphs:
        for placeholder, new_text_val in replacements.items():
            should_be_bold = placeholder in bold_keys
            replace_placeholder_in_paragraph(paragraph, placeholder, new_text_val, is_bold=should_be_bold)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph_in_cell in cell.paragraphs:
                    for placeholder, new_text_val in replacements.items():
                        should_be_bold = placeholder in bold_keys
                        replace_placeholder_in_paragraph(paragraph_in_cell, placeholder, new_text_val, is_bold=should_be_bold)
    return doc

def process_comissao_performance(doc: Document, comissao_existe: bool) -> Document: #
    # ... (código completo da função process_comissao_performance de doc_utils.py)
    if comissao_existe:
        return doc

    body_element = doc.element.body
    if not body_element:
        return doc

    item_6_2_pattern = re.compile(r"^\s*6\.2[\s\.\t]")
    rationale_start_pattern = re.compile(r"^\s*Comissão de Performance", re.IGNORECASE)
    generic_item_6_x_pattern = re.compile(r"^\s*6\.\d+[\s\.\t]")
    item_6_3_pattern = re.compile(r"^\s*6\.3[\s\.\t]")
    sentence_to_remove_in_6_3_base = "As remunerações previstas nas Cláusulas 6.1 e 6.2 acima são denominadas, em conjunto, como “Comissões”"

    elements_to_remove = []
    paragraph_6_3_element_ref = None 
    found_6_2_title_flag = False
    in_rationale_removal_flag = False
    
    child_elements = list(body_element.iterchildren())

    for oxml_el in child_elements:
        if not isinstance(oxml_el, CT_P):
            found_6_2_title_flag = False
            in_rationale_removal_flag = False
            continue

        para_obj = DocxParagraphObject(oxml_el, doc)
        para_text_stripped = para_obj.text.strip()

        if found_6_2_title_flag:
            if rationale_start_pattern.match(para_text_stripped):
                elements_to_remove.append(oxml_el)
                in_rationale_removal_flag = True
                found_6_2_title_flag = False
            elif generic_item_6_x_pattern.match(para_text_stripped):
                found_6_2_title_flag = False
                in_rationale_removal_flag = False
                if item_6_3_pattern.match(para_text_stripped) and paragraph_6_3_element_ref is None:
                    paragraph_6_3_element_ref = oxml_el
            else:
                elements_to_remove.append(oxml_el)
        
        elif in_rationale_removal_flag:
            if generic_item_6_x_pattern.match(para_text_stripped):
                in_rationale_removal_flag = False
                if item_6_3_pattern.match(para_text_stripped) and paragraph_6_3_element_ref is None:
                    paragraph_6_3_element_ref = oxml_el
            else:
                elements_to_remove.append(oxml_el)
        
        elif item_6_2_pattern.match(para_text_stripped):
            elements_to_remove.append(oxml_el)
            found_6_2_title_flag = True
        elif item_6_3_pattern.match(para_text_stripped) and paragraph_6_3_element_ref is None:
            paragraph_6_3_element_ref = oxml_el

    for el_to_remove in elements_to_remove:
        try:
            body_element.remove(el_to_remove)
        except ValueError:
            pass

    if paragraph_6_3_element_ref and isinstance(paragraph_6_3_element_ref, CT_P):
        para_6_3_obj = DocxParagraphObject(paragraph_6_3_element_ref, doc)
        current_text_6_3 = para_6_3_obj.text
        modified_text_6_3 = current_text_6_3
        
        possible_endings = [".”", ".", "”.", "!”", "!"] 
        sentence_found_and_removed = False
        for ending in possible_endings:
            sentence_variant = sentence_to_remove_in_6_3_base[:-1] + ending if sentence_to_remove_in_6_3_base.endswith('”') else sentence_to_remove_in_6_3_base + ending
            if sentence_variant in modified_text_6_3:
                modified_text_6_3 = modified_text_6_3.replace(sentence_variant, "")
                sentence_found_and_removed = True
                break
        if not sentence_found_and_removed and sentence_to_remove_in_6_3_base in modified_text_6_3:
            modified_text_6_3 = modified_text_6_3.replace(sentence_to_remove_in_6_3_base, "")

        modified_text_6_3 = re.sub(r'\s{2,}', ' ', modified_text_6_3).strip()

        for run_idx in range(len(para_6_3_obj.runs) - 1, -1, -1):
            para_6_3_obj._p.remove(para_6_3_obj.runs[run_idx]._r)
        
        if modified_text_6_3:
            new_run = para_6_3_obj.add_run(modified_text_6_3)
            new_run.font.name = 'Calibri'
            new_run.font.size = Pt(11)
            new_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            new_run.font.bold = None
            new_run.font.italic = None

    all_paragraphs_after_changes = doc.paragraphs
    for para_obj in all_paragraphs_after_changes:
        current_para_text = para_obj.text
        renumber_match = re.match(r"^(\s*6\.)(\d+)([\s\.\t])", current_para_text)
        
        if renumber_match:
            captured_prefix_before_num = renumber_match.group(1)
            captured_original_num_str = renumber_match.group(2)
            captured_separator_after_num = renumber_match.group(3)
            text_after_full_prefix = current_para_text[renumber_match.end(0):]

            try:
                original_num_int = int(captured_original_num_str)
                if 3 <= original_num_int <= 10:
                    new_num_int = original_num_int - 1
                    new_full_prefix = captured_prefix_before_num + str(new_num_int) + captured_separator_after_num
                    new_paragraph_text = new_full_prefix + text_after_full_prefix

                    original_font_name = 'Calibri'
                    original_font_size = Pt(11)
                    original_font_color_rgb = RGBColor(0,0,0)
                    original_bold = None
                    
                    if para_obj.runs:
                        first_run = para_obj.runs[0]
                        original_font_name = first_run.font.name if first_run.font.name else 'Calibri'
                        original_font_size = first_run.font.size if first_run.font.size else Pt(11)
                        original_bold = first_run.font.bold
                        if first_run.font.color and first_run.font.color.rgb:
                            original_font_color_rgb = first_run.font.color.rgb

                    for run_idx in range(len(para_obj.runs) - 1, -1, -1):
                        para_obj._p.remove(para_obj.runs[run_idx]._r)
                    
                    new_run = para_obj.add_run(new_paragraph_text)
                    new_run.font.name = original_font_name
                    new_run.font.size = original_font_size
                    new_run.font.color.rgb = original_font_color_rgb
                    new_run.font.bold = original_bold
            except ValueError:
                continue
    return doc


def adjust_anexo_layout(doc: Document) -> Document: #
    # ... (código completo da função adjust_anexo_layout de doc_utils.py)
    anexo_pattern = re.compile(r"^\s*ANEXO\s+([IVXLCDM]+)(\s*-|\s\.|\s|$)", re.IGNORECASE)
    body_element = doc.element.body
    if not body_element:
        return doc

    elements_to_remove = set()
    child_elements = list(body_element.iterchildren())

    for i, current_oxml_el in enumerate(child_elements):
        if isinstance(current_oxml_el, CT_P):
            current_para_obj = DocxParagraphObject(current_oxml_el, doc)
            if anexo_pattern.match(current_para_obj.text.strip()):
                j = i - 1
                while j >= 0:
                    prev_oxml_el_candidate = child_elements[j]
                    if isinstance(prev_oxml_el_candidate, CT_P):
                        prev_para_candidate_obj = DocxParagraphObject(prev_oxml_el_candidate, doc)
                        if not prev_para_candidate_obj.text.strip():
                            elements_to_remove.add(prev_oxml_el_candidate)
                            j -= 1
                        else:
                            break 
                    else:
                        break
    
    if elements_to_remove:
        for el_to_remove in elements_to_remove:
            try:
                body_element.remove(el_to_remove)
            except ValueError:
                pass

    current_paragraphs_after_cleaning = doc.paragraphs
    for idx, para_obj in enumerate(current_paragraphs_after_cleaning):
        if anexo_pattern.match(para_obj.text.strip()):
            if idx == 0:
                continue
            para_obj.paragraph_format.page_break_before = True
    return doc


def generate_docx_dcm(template_path_or_file, data_replacements, bold_placeholders, comissao_existe): #
    document = Document(template_path_or_file)
    filled_document = replace_placeholders_in_document(document, data_replacements, bold_placeholders)
    filled_document = process_comissao_performance(filled_document, comissao_existe)
    filled_document = adjust_anexo_layout(filled_document)

    doc_io = BytesIO()
    filled_document.save(doc_io)
    doc_io.seek(0)
    return doc_io

def generate_docx_coordenacao(template_path_or_file, data_replacements):
    document = Document(template_path_or_file)
    # Para coordenação, não estamos usando bold_keys específicos inicialmente,
    # a formatação é aplicada diretamente pela replace_placeholder_in_paragraph (Calibri 11pt).
    filled_document = replace_placeholders_in_document(document, data_replacements, bold_keys=set())

    doc_io = BytesIO()
    filled_document.save(doc_io)
    doc_io.seek(0)
    return doc_io

# --- Lógica de Preparação de Dados para "Coordenação" ---
COORD_FIELD_CONFIGS = {
    "CRI": [
        ("Emissora:", "coord_emissora", "Empresa Emissora Exemplo S.A."),
        ("CNPJ Emissora:", "coord_cnpj_emissora", "00.000.000/0001-00"),
        ("Cópia (Nome):", "coord_copia_nome", ""),
        ("Email Cópia:", "coord_copia_email", ""),
        ("Valor Total:", "coord_valor_total_str", "100000000,00"), #
        ("Remuneração:", "coord_remuneracao_str", "60000,00"), #
        ("Prazo:", "coord_prazo", "10 anos"), #
        ("Lastro:", "coord_lastro", "Ex. (CCI)"), #
        ("Devedora:", "coord_devedora", "Empresa Devedora Exemplo Ltda."), #
        ("Destinação:", "coord_destinacao", "Alocação dos recursos arrecadados com a emissão dos CRI."), #
        ("Remuneração Título:", "coord_remuneracao_titulo", "Ex. (CDI + 6%)"), #
        ("Amortização Principal:", "coord_amortizacao_principal", "Ex. (Mensal)"), #
        ("Pagamento Juros:", "coord_pagamento_juros", "Ex. (Mensal)"), #
        ("Garantias:", "coord_garantias", "1) Exemplo de Garantia 1\n2) Exemplo de Garantia 2"), #
        ("Assessor Legal:", "coord_assessor_legal", "Escritório de Advocacia Exemplo") #
    ],
    "CRA": [
        ("Emissora:", "coord_emissora", "Empresa Emissora Exemplo S.A."),
        ("CNPJ Emissora:", "coord_cnpj_emissora", "00.000.000/0001-00"),
        ("Cópia (Nome):", "coord_copia_nome", ""),
        ("Email Cópia:", "coord_copia_email", ""),
        ("Valor Total:", "coord_valor_total_str", "50000000,00"), #
        ("Remuneração:", "coord_remuneracao_str", "50000,00"), #
        ("Prazo:", "coord_prazo", "5 anos"), #
        ("Lastro:", "coord_lastro", "Ex. (CDA/WA)"), #
        ("Devedora:", "coord_devedora", "Agropecuária Exemplo S.A."), #
        ("Destinação:", "coord_destinacao", "Alocação dos recursos arrecadados com a emissão dos CRA."), #
        ("Remuneração Título:", "coord_remuneracao_titulo", "Ex. (IPCA + 7%)"), #
        ("Amortização Principal:", "coord_amortizacao_principal", "Ex. (Anual)"), #
        ("Pagamento Juros:", "coord_pagamento_juros", "Ex. (Semestral)"), #
        ("Garantias:", "coord_garantias", "1) Penhor Agrícola\n2) Fiança"), #
        ("Assessor Legal:", "coord_assessor_legal", "Consultoria Jurídica Exemplo") #
    ],
    "Debênture": [
        ("Emissora:", "coord_emissora", "Empresa Emissora Debêntures S.A."),
        ("CNPJ Emissora:", "coord_cnpj_emissora", "11.111.111/0001-11"),
        ("Cópia (Nome):", "coord_copia_nome", ""),
        ("Email Cópia:", "coord_copia_email", ""),
        ("Valor Total:", "coord_valor_total_str", "200000000,00"), #
        ("Remuneração:", "coord_remuneracao_str", "120000,00"), #
        ("Prazo:", "coord_prazo", "7 anos"), #
        ("Remuneração Título:", "coord_remuneracao_titulo", "Ex. (NTN-B + 1,5%)"), #
        ("Amortização Principal:", "coord_amortizacao_principal", "Ex. (Ao final)"), #
        ("Pagamento Juros:", "coord_pagamento_juros", "Ex. (Anual)"), #
        ("Garantias:", "coord_garantias", "1) Fiança Bancária\n2) Alienação Fiduciária de Ações") #
        # Debênture não possui Lastro, Devedora, Destinação (explícita como nos outros), Assessor Legal nos scripts originais
    ]
}


def prepare_coordenacao_data(inputs: dict, offer_type: str) -> tuple[dict, list]:
    """
    Prepara os dados para o template do documento de Coordenação.
    Retorna um dicionário de substituições e uma lista de erros.
    """
    errors = []
    data_to_replace = {}

    try:
        valor_total_str = inputs.get("coord_valor_total_str", "0").replace(",", ".") #
        valor_total = float(valor_total_str)
    except ValueError:
        errors.append(f"Valor Total inválido: '{inputs.get('coord_valor_total_str', '')}'. Use formato numérico como 100000,00.")
        valor_total = 0.0

    try:
        remuneracao_str = inputs.get("coord_remuneracao_str", "0").replace(",", ".") #
        remuneracao_valor = float(remuneracao_str) # Nos scripts de coordenação, a remuneração é um valor, não percentual.
    except ValueError:
        errors.append(f"Remuneração inválida: '{inputs.get('coord_remuneracao_str', '')}'. Use formato numérico como 60000,00.")
        remuneracao_valor = 0.0

    if errors:
        return {}, errors

    data_hoje = datetime.now()
    dia = data_hoje.day
    meses_pt = ["janeiro", "fevereiro", "março", "abril", "maio", "junho",
                "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    mes = meses_pt[data_hoje.month - 1]
    ano = data_hoje.year # Adicionado para ter o ano

    valor_total_ext = num2words(valor_total, lang='pt_BR', to='currency') #
    
    qtd_total = valor_total / 1000
    qtd_total_ext = num2words(qtd_total, lang='pt_BR')

    remuneracao_ext = num2words(remuneracao_valor, lang='pt_BR', to='currency') #

    data_1ano_obj = data_hoje + relativedelta(years=1) #
    data_1ano_ext = f'{data_1ano_obj.day} de {meses_pt[data_1ano_obj.month - 1]} de {data_1ano_obj.year}' #

    data_20dias_obj = data_hoje + timedelta(days=20) #
    data_20dias_ext = f'{data_20dias_obj.day} de {meses_pt[data_20dias_obj.month - 1]} de {data_20dias_obj.year}' #

    # Formatando valores para o padrão brasileiro (xxx.xxx,xx)
    valor_total_fmt = f"{valor_total:_.2f}".replace("_", "X").replace(".", ",").replace("X", ".") #
    remuneracao_fmt = f"{remuneracao_valor:_.2f}".replace("_", "X").replace(".", ",").replace("X", ".") #
    qtd_total_fmt = f"{qtd_total:_.0f}".replace("_", "X").replace(".", ",").replace("X", ".")


    data_to_replace = {
        "[[Dia]]": str(dia), #
        "[[Mes]]": str(mes), #
        "[[Ano]]": str(ano), # Adicionado
        "[[Emissora]]": inputs.get("coord_emissora", ""), #
        "[[CNPJ_Emissora]]": inputs.get("coord_cnpj_emissora", ""), #
        "[[Copia]]": inputs.get("coord_copia_nome", ""), #
        "[[Email_Copia]]": inputs.get("coord_copia_email", ""), #
        "[[Valor_Total]]": valor_total_fmt, #
        "[[Valor_Total_Ext]]": valor_total_ext.capitalize(), #
        "[[Qtd_Total]]": qtd_total_fmt, # Se for usar
        "[[Qtd_Total_Ext]]": qtd_total_ext, # Se for usar
        "[[Remuneracao]]": remuneracao_fmt, #
        "[[Remuneracao_Ext]]": remuneracao_ext.capitalize(), #
        "[[Data_1ano]]": data_1ano_ext, #
        "[[Data_20dias]]": data_20dias_ext, #
        "[[Prazo]]": inputs.get("coord_prazo", ""), #
        "[[Remuneracao_Titulo]]": inputs.get("coord_remuneracao_titulo", ""), #
        "[[Amor_Princ]]": inputs.get("coord_amortizacao_principal", ""), #
        "[[Pgto_Juros]]": inputs.get("coord_pagamento_juros", ""), #
        "[[Garantias]]": inputs.get("coord_garantias", "") #
    }

    if offer_type in ["CRI", "CRA"]:
        data_to_replace["[[Lastro]]"] = inputs.get("coord_lastro", "") #
        data_to_replace["[[Devedora]]"] = inputs.get("coord_devedora", "") #
        data_to_replace["[[Destinacao]]"] = inputs.get("coord_destinacao", "") #
        data_to_replace["[[Assessor]]"] = inputs.get("coord_assessor_legal", "") # (Placeholder [[Assessor]] usado em cri.py/cra.py)
    
    # Adicionar quaisquer outros placeholders específicos se necessário

    return data_to_replace, errors