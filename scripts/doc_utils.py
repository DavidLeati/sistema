import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph as DocxParagraphObject
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.oxml.table import CT_Tbl # Não usado diretamente, mas bom manter se houver manipulação de tabelas

def replace_placeholder_in_paragraph(paragraph: DocxParagraph, placeholder: str, new_text_val: str, is_bold: bool = False):
    """
    Substitui um placeholder em um parágrafo, aplicando negrito ao texto substituído se is_bold=True.
    Esta versão reconstrói o parágrafo para aplicar negrito de forma mais precisa.
    """
    str_new_text_val = str(new_text_val)

    current_paragraph_text = "".join(run.text for run in paragraph.runs)

    if placeholder not in current_paragraph_text:
        return

    temp_paragraph_text_parts = []
    last_idx = 0
    while placeholder in current_paragraph_text[last_idx:]:
        start_idx = current_paragraph_text.find(placeholder, last_idx)
        if start_idx == -1:
            break
        
        if start_idx > last_idx:
            temp_paragraph_text_parts.append( (current_paragraph_text[last_idx:start_idx], False) )
        
        temp_paragraph_text_parts.append( (str_new_text_val, is_bold) )
        last_idx = start_idx + len(placeholder)

    if last_idx < len(current_paragraph_text):
        temp_paragraph_text_parts.append( (current_paragraph_text[last_idx:], False) )

    # Limpa runs existentes antes de adicionar novas
    # Iterar para trás para remover é mais seguro se a remoção afetar índices
    for _ in range(len(paragraph.runs)):
        if paragraph.runs: # Verifica se ainda há runs
            p_element = paragraph._p
            p_element.remove(paragraph.runs[0]._r) # Remove a primeira run

    for text_content, should_be_bold_segment in temp_paragraph_text_parts:
        if not text_content:
            continue
        run = paragraph.add_run(text_content)
        # Preservar a formatação original do parágrafo, se possível, ou definir um padrão.
        # Aqui, definimos um padrão. Se precisar herdar, a lógica seria mais complexa.
        run.font.name = 'Calibri' # Padrão
        run.font.size = Pt(11)    # Padrão
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00) # Padrão (preto)
        run.font.bold = should_be_bold_segment
        # run.font.italic = False # Resetar outros estilos se necessário


def replace_placeholders_in_document(doc: Document, replacements: dict, bold_keys: set) -> Document:
    for paragraph in doc.paragraphs:
        for placeholder, new_text_val in replacements.items():
            should_be_bold = placeholder in bold_keys
            replace_placeholder_in_paragraph(paragraph, placeholder, new_text_val, is_bold=should_be_bold)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph_in_cell in cell.paragraphs:
                    for placeholder, new_text_val in replacements.items():
                        should_be_bold = placeholder in bold_keys
                        replace_placeholder_in_paragraph(paragraph_in_cell, placeholder, new_text_val, is_bold=should_be_bold)
    return doc

def process_comissao_performance(doc: Document, comissao_existe: bool) -> Document:
    if comissao_existe:
        return doc

    body_element = doc.element.body
    if not body_element:
        return doc

    item_6_2_pattern = re.compile(r"^\s*6\.2[\s\.\t]")
    rationale_start_pattern = re.compile(r"^\s*Comissão de Performance", re.IGNORECASE)
    generic_item_6_x_pattern = re.compile(r"^\s*6\.\d+[\s\.\t]")
    item_6_3_pattern = re.compile(r"^\s*6\.3[\s\.\t]")
    sentence_to_remove_in_6_3_base = "As remunerações previstas nas Cláusulas 6.1 e 6.2 acima são denominadas, em conjunto, como “Comissões”"

    elements_to_remove = []
    paragraph_6_3_element_ref = None 
    found_6_2_title_flag = False
    in_rationale_removal_flag = False
    
    child_elements = list(body_element.iterchildren())

    for oxml_el in child_elements:
        if not isinstance(oxml_el, CT_P):
            found_6_2_title_flag = False
            in_rationale_removal_flag = False
            continue

        para_obj = DocxParagraphObject(oxml_el, doc) # Use a alias correta
        para_text_stripped = para_obj.text.strip()

        if found_6_2_title_flag:
            if rationale_start_pattern.match(para_text_stripped):
                elements_to_remove.append(oxml_el)
                in_rationale_removal_flag = True
                found_6_2_title_flag = False
            elif generic_item_6_x_pattern.match(para_text_stripped):
                found_6_2_title_flag = False
                in_rationale_removal_flag = False
                if item_6_3_pattern.match(para_text_stripped) and paragraph_6_3_element_ref is None:
                    paragraph_6_3_element_ref = oxml_el
            else:
                elements_to_remove.append(oxml_el)
        
        elif in_rationale_removal_flag:
            if generic_item_6_x_pattern.match(para_text_stripped):
                in_rationale_removal_flag = False
                if item_6_3_pattern.match(para_text_stripped) and paragraph_6_3_element_ref is None:
                    paragraph_6_3_element_ref = oxml_el
            else:
                elements_to_remove.append(oxml_el)
        
        elif item_6_2_pattern.match(para_text_stripped):
            elements_to_remove.append(oxml_el)
            found_6_2_title_flag = True
        elif item_6_3_pattern.match(para_text_stripped) and paragraph_6_3_element_ref is None:
            paragraph_6_3_element_ref = oxml_el

    for el_to_remove in elements_to_remove:
        try:
            body_element.remove(el_to_remove)
        except ValueError:
            pass # Elemento pode já ter sido removido ou não ser filho direto

    if paragraph_6_3_element_ref and isinstance(paragraph_6_3_element_ref, CT_P):
        para_6_3_obj = DocxParagraphObject(paragraph_6_3_element_ref, doc) # Use a alias correta
        current_text_6_3 = para_6_3_obj.text
        modified_text_6_3 = current_text_6_3
        
        possible_endings = [".”", ".", "”.", "!”", "!"] 
        sentence_found_and_removed = False
        for ending in possible_endings:
            sentence_variant = sentence_to_remove_in_6_3_base[:-1] + ending if sentence_to_remove_in_6_3_base.endswith('”') else sentence_to_remove_in_6_3_base + ending
            if sentence_variant in modified_text_6_3:
                modified_text_6_3 = modified_text_6_3.replace(sentence_variant, "")
                sentence_found_and_removed = True
                break
        if not sentence_found_and_removed and sentence_to_remove_in_6_3_base in modified_text_6_3: # Fallback
            modified_text_6_3 = modified_text_6_3.replace(sentence_to_remove_in_6_3_base, "")

        modified_text_6_3 = re.sub(r'\s{2,}', ' ', modified_text_6_3).strip()

        for run_idx in range(len(para_6_3_obj.runs) - 1, -1, -1):
            para_6_3_obj._p.remove(para_6_3_obj.runs[run_idx]._r)
        
        if modified_text_6_3: # Só adiciona run se houver texto
            new_run = para_6_3_obj.add_run(modified_text_6_3)
            # Aplicar formatação padrão (Calibri, 11, Preto)
            new_run.font.name = 'Calibri'
            new_run.font.size = Pt(11)
            new_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            # Resetar outros estilos se necessário (bold, italic etc.)
            new_run.font.bold = None # Ou False, dependendo do desejado
            new_run.font.italic = None # Ou False

    # Renumerar itens
    all_paragraphs_after_changes = doc.paragraphs
    for para_obj in all_paragraphs_after_changes:
        current_para_text = para_obj.text
        renumber_match = re.match(r"^(\s*6\.)(\d+)([\s\.\t])", current_para_text)
        
        if renumber_match:
            captured_prefix_before_num = renumber_match.group(1)
            captured_original_num_str = renumber_match.group(2)
            captured_separator_after_num = renumber_match.group(3)
            text_after_full_prefix = current_para_text[renumber_match.end(0):]

            try:
                original_num_int = int(captured_original_num_str)
                if 3 <= original_num_int <= 10: # Itens originais 6.3 a 6.10
                    new_num_int = original_num_int - 1
                    new_full_prefix = captured_prefix_before_num + str(new_num_int) + captured_separator_after_num
                    new_paragraph_text = new_full_prefix + text_after_full_prefix

                    # Preservar a formatação do primeiro run, se existir, e aplicá-la.
                    # Ou, mais simplesmente, aplicar uma formatação padrão.
                    # Vamos aplicar uma formatação padrão para consistência.
                    original_font_name = 'Calibri'
                    original_font_size = Pt(11)
                    original_font_color_rgb = RGBColor(0,0,0)
                    original_bold = None # Default
                    
                    if para_obj.runs: # Tenta pegar do primeiro run existente
                        first_run = para_obj.runs[0]
                        original_font_name = first_run.font.name if first_run.font.name else 'Calibri'
                        original_font_size = first_run.font.size if first_run.font.size else Pt(11)
                        original_bold = first_run.font.bold
                        if first_run.font.color and first_run.font.color.rgb:
                            original_font_color_rgb = first_run.font.color.rgb
                        # else: mantém o default RGBColor(0,0,0)

                    # Limpa runs existentes
                    for run_idx in range(len(para_obj.runs) - 1, -1, -1):
                        para_obj._p.remove(para_obj.runs[run_idx]._r)
                    
                    new_run = para_obj.add_run(new_paragraph_text)
                    new_run.font.name = original_font_name
                    new_run.font.size = original_font_size
                    new_run.font.color.rgb = original_font_color_rgb
                    new_run.font.bold = original_bold # Aplica o negrito original
                    # Resetar outros se necessário:
                    # new_run.font.italic = None 
                    # new_run.font.underline = None

            except ValueError:
                continue # Se captured_original_num_str não for um int válido
    return doc

def adjust_anexo_layout(doc: Document) -> Document:
    anexo_pattern = re.compile(r"^\s*ANEXO\s+([IVXLCDM]+)(\s*-|\s\.|\s|$)", re.IGNORECASE)
    body_element = doc.element.body
    if not body_element:
        return doc

    elements_to_remove = set()
    child_elements = list(body_element.iterchildren())

    for i, current_oxml_el in enumerate(child_elements):
        if isinstance(current_oxml_el, CT_P):
            current_para_obj = DocxParagraphObject(current_oxml_el, doc)
            if anexo_pattern.match(current_para_obj.text.strip()):
                j = i - 1
                while j >= 0:
                    prev_oxml_el_candidate = child_elements[j]
                    if isinstance(prev_oxml_el_candidate, CT_P):
                        prev_para_candidate_obj = DocxParagraphObject(prev_oxml_el_candidate, doc)
                        if not prev_para_candidate_obj.text.strip():
                            elements_to_remove.add(prev_oxml_el_candidate)
                            j -= 1
                        else:
                            break 
                    else: # Não é parágrafo (ex: tabela), para.
                        break
    
    if elements_to_remove:
        for el_to_remove in elements_to_remove:
            try:
                body_element.remove(el_to_remove)
            except ValueError:
                pass

    current_paragraphs_after_cleaning = doc.paragraphs
    for idx, para_obj in enumerate(current_paragraphs_after_cleaning):
        if anexo_pattern.match(para_obj.text.strip()):
            if idx == 0: # Não adiciona quebra antes do primeiro parágrafo absoluto
                continue
            para_obj.paragraph_format.page_break_before = True
    return doc

def generate_docx(template_path_or_file, data_replacements, bold_placeholders, comissao_existe):
    """
    Gera o documento DOCX preenchido.
    """
    document = Document(template_path_or_file)
    filled_document = replace_placeholders_in_document(document, data_replacements, bold_placeholders)
    filled_document = process_comissao_performance(filled_document, comissao_existe)
    filled_document = adjust_anexo_layout(filled_document)

    doc_io = BytesIO()
    filled_document.save(doc_io)
    doc_io.seek(0)
    return doc_io