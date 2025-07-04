# proposal_system/document_processing/docx_utils.py
import re
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph as DocxParagraphObject
from docx.oxml.text.paragraph import CT_P

def find_next_placeholder_details(text_to_search: str, placeholders_map: dict, search_start_index: int = 0): #
    first_placeholder_found = None #
    earliest_position = -1 #

    for ph_key in placeholders_map.keys(): #
        position = text_to_search.find(ph_key, search_start_index) #
        if position != -1: #
            if earliest_position == -1 or position < earliest_position: #
                earliest_position = position #
                first_placeholder_found = ph_key #
    
    return first_placeholder_found, earliest_position #

def replace_placeholders_in_paragraph(paragraph: DocxParagraphObject, 
                                                replacements: dict, 
                                                bold_placeholders: set): #
    current_paragraph_text = "".join(run.text for run in paragraph.runs) #

    if not any(ph in current_paragraph_text for ph in replacements.keys()): #
        return #

    text_segments = []   #
    current_search_idx = 0 #

    while current_search_idx < len(current_paragraph_text): #
        placeholder_key, placeholder_start_idx = find_next_placeholder_details( #
            current_paragraph_text, replacements, current_search_idx #
        )

        if placeholder_key is None: #
            if current_search_idx < len(current_paragraph_text): #
                text_segments.append((current_paragraph_text[current_search_idx:], False)) #
            break   #

        if placeholder_start_idx > current_search_idx: #
            text_segments.append((current_paragraph_text[current_search_idx:placeholder_start_idx], False)) #

        replacement_text = str(replacements.get(placeholder_key, placeholder_key))  #
        should_be_bold = placeholder_key in bold_placeholders #
        text_segments.append((replacement_text, should_be_bold)) #

        current_search_idx = placeholder_start_idx + len(placeholder_key) #

    while paragraph.runs: #
        p_element = paragraph._p   #
        r_element = paragraph.runs[0]._r   #
        p_element.remove(r_element) #

    for text_content, should_be_bold_segment in text_segments: #
        if not text_content:   #
            continue #
        run = paragraph.add_run(text_content) #
        run.font.name = 'Calibri' #
        run.font.size = Pt(11) #
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00) #
        run.font.bold = should_be_bold_segment #

def replace_placeholders_in_document(doc, replacements_map: dict, placeholders_to_make_bold: set = None): #
    if placeholders_to_make_bold is None: #
        placeholders_to_make_bold = set() #

    for paragraph in doc.paragraphs: #
        replace_placeholders_in_paragraph(paragraph, replacements_map, placeholders_to_make_bold) #
    
    for table in doc.tables: #
        for row in table.rows: #
            for cell in row.cells: #
                for paragraph_in_cell in cell.paragraphs: #
                    replace_placeholders_in_paragraph(paragraph_in_cell, replacements_map, placeholders_to_make_bold) #
    return doc #

def process_comissao_performance(doc, comissao_existe: bool): #
    if comissao_existe: #
        return doc #

    body_element = doc.element.body #
    if not body_element: #
        return doc #

    item_6_2_pattern = re.compile(r"^\s*6\.2[\s\.\t]") #
    rationale_start_pattern = re.compile(r"^\s*Comissão de Performance", re.IGNORECASE) #
    generic_item_6_x_pattern = re.compile(r"^\s*6\.\d+[\s\.\t]") #
    item_6_3_pattern = re.compile(r"^\s*6\.3[\s\.\t]") #
    sentence_to_remove_in_6_3_base = "As remunerações previstas nas Cláusulas 6.1 e 6.2 acima são denominadas, em conjunto, como “Comissões”" #

    elements_to_remove = [] #
    paragraph_6_3_element_ref = None  #
    found_6_2_title_flag = False #
    in_rationale_removal_flag = False #
    
    child_elements = list(body_element.iterchildren()) #

    for oxml_el in child_elements: #
        if not isinstance(oxml_el, CT_P): #
            found_6_2_title_flag = False #
            in_rationale_removal_flag = False #
            continue #

        para_obj = DocxParagraphObject(oxml_el, doc) #
        para_text_stripped = para_obj.text.strip() #

        if found_6_2_title_flag: #
            if rationale_start_pattern.match(para_text_stripped): #
                elements_to_remove.append(oxml_el) #
                in_rationale_removal_flag = True #
                found_6_2_title_flag = False #
            elif generic_item_6_x_pattern.match(para_text_stripped): #
                found_6_2_title_flag = False #
                in_rationale_removal_flag = False #
                if item_6_3_pattern.match(para_text_stripped) and paragraph_6_3_element_ref is None: #
                    paragraph_6_3_element_ref = oxml_el #
            else:
                elements_to_remove.append(oxml_el) #
        
        elif in_rationale_removal_flag: #
            if generic_item_6_x_pattern.match(para_text_stripped): #
                in_rationale_removal_flag = False #
                if item_6_3_pattern.match(para_text_stripped) and paragraph_6_3_element_ref is None: #
                    paragraph_6_3_element_ref = oxml_el #
            else:
                elements_to_remove.append(oxml_el) #
        
        elif item_6_2_pattern.match(para_text_stripped): #
            elements_to_remove.append(oxml_el) #
            found_6_2_title_flag = True #
        elif item_6_3_pattern.match(para_text_stripped) and paragraph_6_3_element_ref is None: #
            paragraph_6_3_element_ref = oxml_el #

    for el_to_remove in elements_to_remove: #
        try:
            body_element.remove(el_to_remove) #
        except ValueError: #
            pass #

    if paragraph_6_3_element_ref and isinstance(paragraph_6_3_element_ref, CT_P): #
        para_6_3_obj = DocxParagraphObject(paragraph_6_3_element_ref, doc) #
        current_text_6_3 = para_6_3_obj.text #
        modified_text_6_3 = current_text_6_3 #
        
        possible_endings = [".”", ".", "”.", "!”", "!"]  #
        sentence_found_and_removed = False #
        for ending in possible_endings: #
            sentence_variant = sentence_to_remove_in_6_3_base[:-1] + ending if sentence_to_remove_in_6_3_base.endswith('”') else sentence_to_remove_in_6_3_base + ending #
            if sentence_variant in modified_text_6_3: #
                modified_text_6_3 = modified_text_6_3.replace(sentence_variant, "") #
                sentence_found_and_removed = True #
                break #
        if not sentence_found_and_removed and sentence_to_remove_in_6_3_base in modified_text_6_3: #
            modified_text_6_3 = modified_text_6_3.replace(sentence_to_remove_in_6_3_base, "") #

        modified_text_6_3 = re.sub(r'\s{2,}', ' ', modified_text_6_3).strip() #

        for run_idx in range(len(para_6_3_obj.runs) - 1, -1, -1): #
            para_6_3_obj._p.remove(para_6_3_obj.runs[run_idx]._r) #
        
        if modified_text_6_3: #
            new_run = para_6_3_obj.add_run(modified_text_6_3) #
            new_run.font.name = 'Calibri' #
            new_run.font.size = Pt(11) #
            new_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00) #
            new_run.font.bold = None #
            new_run.font.italic = None #
        
    return doc #


def adjust_anexo_layout(doc): #
    anexo_pattern = re.compile(r"^\s*ANEXO\s+([IVXLCDM]+)(\s*-|\s\.|\s|$)", re.IGNORECASE) #
    body_element = doc.element.body #
    if not body_element: #
        return doc #

    elements_to_remove = set() #
    child_elements = list(body_element.iterchildren()) #

    for i, current_oxml_el in enumerate(child_elements): #
        if isinstance(current_oxml_el, CT_P): #
            current_para_obj = DocxParagraphObject(current_oxml_el, doc) #
            if anexo_pattern.match(current_para_obj.text.strip()): #
                j = i - 1 #
                while j >= 0: #
                    prev_oxml_el_candidate = child_elements[j] #
                    if isinstance(prev_oxml_el_candidate, CT_P): #
                        prev_para_candidate_obj = DocxParagraphObject(prev_oxml_el_candidate, doc) #
                        if not prev_para_candidate_obj.text.strip(): #
                            elements_to_remove.add(prev_oxml_el_candidate) #
                            j -= 1 #
                        else:
                            break  #
                    else:
                        break #
    
    if elements_to_remove: #
        for el_to_remove in elements_to_remove: #
            try:
                body_element.remove(el_to_remove) #
            except ValueError: #
                pass #

    current_paragraphs_after_cleaning = doc.paragraphs #
    for idx, para_obj in enumerate(current_paragraphs_after_cleaning): #
        if anexo_pattern.match(para_obj.text.strip()): #
            if idx == 0: #
                continue #
            para_obj.paragraph_format.page_break_before = True #
    return doc #

def rescan_and_renumber_document(doc):
    """
    Verifica todo o documento e reordena os itens de listas numéricas
    que possam estar fora de sequência. Garante a fonte Calibri 11pt.
    """
    anexo_pattern = re.compile(r"^\s*ANEXO\s+[IVXLCDM]+", re.IGNORECASE)
    numeric_item_pattern = re.compile(r"^\s*(\d+(?:\.\d+)*)[\.\s\t]")
    counters = {}
    
    for para in doc.paragraphs:
        # Usaremos o texto do primeiro "run" para a verificação, é mais seguro.
        if not para.runs:
            continue
        
        if anexo_pattern.match(para.text):
            break

        # Concatena o texto dos runs para ter o texto completo do parágrafo
        para_text = para.text
        match = numeric_item_pattern.match(para_text)

        if not match:
            continue

        current_number_str = match.group(1)
        level = current_number_str.count('.')
        
        for i in range(level + 1, len(counters)):
            counters[i] = 0
            
        counters[level] = counters.get(level, 0) + 1
        
        expected_number_parts = [str(counters.get(i, 1)) for i in range(level + 1)]
        expected_number_str = ".".join(expected_number_parts)

        if current_number_str != expected_number_str:
            # Pega o estilo do primeiro run como base
            original_run = para.runs[0]
            is_bold = original_run.bold
            
            # Recria o texto do parágrafo com o número correto
            new_text = para_text.replace(current_number_str, expected_number_str, 1)
            
            # Limpa os runs antigos do parágrafo
            for run in para.runs:
                para._p.remove(run._r)
            
            # Adiciona um novo run com o texto corrigido e a formatação desejada
            new_run = para.add_run(new_text)
            new_run.font.name = 'Calibri'  # <-- FORÇA A FONTE CORRETA
            new_run.font.size = Pt(11)     # <-- FORÇA O TAMANHO CORRETO
            new_run.bold = is_bold         # <-- Mantém o negrito original
            new_run.font.color.rgb = RGBColor(0, 0, 0)
            
    return doc

def delete_empty_table_rows(doc):
    """
    Percorre todas as tabelas do documento e remove as linhas em que a primeira
    ou a segunda coluna estejam vazias (sem texto).
    Aplica a lógica apenas em tabelas com 2 ou mais colunas.
    """
    # Itera sobre cada tabela no documento
    for table in doc.tables:
        # Pula para a próxima tabela se ela não tiver linhas ou se tiver menos de 2 colunas.
        # Isso previne erros e impede que a lógica afete tabelas de 1 coluna.
        if not table.rows or len(table.columns) < 2:
            continue
        if len(table.columns) == 3:
            continue

        # Cria uma lista para marcar as linhas que devem ser removidas.
        rows_to_delete = []
        
        for row in table.rows:
            # O `if len(row.cells) < 2:` anterior foi removido pois a verificação agora é feita na tabela.
            
            # Pega o texto da primeira e da segunda célula, removendo espaços em branco
            cell_1_text = row.cells[0].text.strip()
            cell_2_text = row.cells[1].text.strip()
            
            # Se qualquer uma das duas células estiver vazia, marca a linha para exclusão
            if not cell_1_text or not cell_2_text:
                rows_to_delete.append(row)

        # Após identificar todas as linhas, remove cada uma delas da tabela
        for row_to_delete in rows_to_delete:
            table._tbl.remove(row_to_delete._tr)
            
    return doc

def process_paragraph_by_marker(doc, marker, keep_paragraph=True):
    """
    Encontra um parágrafo por um marcador de texto único.

    - Se keep_paragraph for True, mantém o parágrafo mas remove o marcador.
    - Se keep_paragraph for False, remove o parágrafo inteiro e a linha
      em branco seguinte, se houver.
    """
    paragraphs_to_delete = []
    
    # Criamos uma lista estática de parágrafos para iterar com segurança
    all_paragraphs = list(doc.paragraphs)
    
    for i, p in enumerate(all_paragraphs):
        if marker in p.text:
            if keep_paragraph:
                # Se for para manter, apenas removemos o marcador do texto
                for run in p.runs:
                    if marker in run.text:
                        run.text = run.text.replace(marker, "")
            else:
                # Se for para apagar, adicionamos o parágrafo à lista de exclusão
                paragraphs_to_delete.append(p)
                
                # Verifica se o parágrafo seguinte existe e está em branco
                if (i + 1) < len(all_paragraphs):
                    next_paragraph = all_paragraphs[i + 1]
                    # Se o parágrafo seguinte não tiver texto (apenas espaços), ele também é marcado para exclusão.
                    if not next_paragraph.text.strip():
                        paragraphs_to_delete.append(next_paragraph)

            # Assume que o marcador é único e para a busca
            break 
            
    # Remove todos os parágrafos marcados para exclusão
    for p_to_delete in paragraphs_to_delete:
        p_element = p_to_delete._element
        if p_element.getparent() is not None:
            p_element.getparent().remove(p_element)
            
    return doc

def apply_keep_with_next_to_headings(doc):
    """
    Aplica a formatação "Manter com o próximo" a todos os parágrafos
    que são identificados como títulos numerados (ex: 1., 1.1., etc.).
    Isso impede que um título fique isolado no final de uma página.
    """
    # Regex para identificar um item de lista numérica (tópicos e subtópicos)
    # O mesmo padrão usado na função de renumeração.
    heading_pattern = re.compile(r"^\s*(\d+(?:\.\d+)*)[\.\s\t]")

    for para in doc.paragraphs:
        if heading_pattern.match(para.text.strip()):
            para.paragraph_format.keep_with_next = True
    
    return doc